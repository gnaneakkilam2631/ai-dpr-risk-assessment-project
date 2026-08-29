import re
import shutil

from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import (
    APIRouter,
    Depends,
    File,
    HTTPException,
    UploadFile,
)

from sqlalchemy.orm import Session

from app.database import get_db
from app.models import Document, Project


# ============================================================
# ROUTER
# ============================================================

router = APIRouter(
    prefix="/documents",
    tags=["Documents"],
)


# ============================================================
# DIRECTORIES
# ============================================================

BASE_DIR = Path(__file__).resolve().parents[2]

UPLOAD_DIR = BASE_DIR / "uploads"

UPLOAD_DIR.mkdir(
    parents=True,
    exist_ok=True,
)


# ============================================================
# ALLOWED FILE TYPES
# ============================================================

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}


# ============================================================
# BASIC HELPERS
# ============================================================

def get_file_extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def get_safe_filename(filename: str) -> str:
    return Path(filename).name.strip()


# ============================================================
# DOCUMENT TEXT EXTRACTION
# ============================================================

def extract_document_text(file_path: Path) -> str:

    extension = file_path.suffix.lower()

    # ========================================================
    # TXT
    # ========================================================

    if extension == ".txt":

        try:
            return file_path.read_text(
                encoding="utf-8"
            )

        except UnicodeDecodeError:
            return file_path.read_text(
                encoding="latin-1"
            )


    # ========================================================
    # PDF
    # ========================================================

    if extension == ".pdf":

        try:
            from pypdf import PdfReader

        except ImportError as exc:

            raise HTTPException(
                status_code=500,
                detail=(
                    "pypdf is not installed. "
                    "Run: pip install pypdf"
                ),
            ) from exc

        try:

            reader = PdfReader(
                str(file_path)
            )

            pages: List[str] = []

            for page in reader.pages:

                page_text = (
                    page.extract_text()
                    or ""
                )

                if page_text.strip():

                    pages.append(
                        page_text
                    )

            return "\n\n".join(
                pages
            )

        except Exception as exc:

            print(
                "[PDF] TEXT EXTRACTION ERROR:",
                exc,
            )

            raise HTTPException(
                status_code=500,
                detail="Could not extract PDF text.",
            ) from exc


    # ========================================================
    # DOCX
    # ========================================================

    if extension == ".docx":

        try:
            from docx import Document as DocxDocument

        except ImportError as exc:

            raise HTTPException(
                status_code=500,
                detail=(
                    "python-docx is not installed. "
                    "Run: pip install python-docx"
                ),
            ) from exc

        try:

            doc = DocxDocument(
                str(file_path)
            )

            paragraphs: List[str] = []

            for paragraph in doc.paragraphs:

                text = paragraph.text.strip()

                if text:
                    paragraphs.append(
                        text
                    )

            return "\n\n".join(
                paragraphs
            )

        except Exception as exc:

            print(
                "[DOCX] TEXT EXTRACTION ERROR:",
                exc,
            )

            raise HTTPException(
                status_code=500,
                detail="Could not extract DOCX text.",
            ) from exc


    raise HTTPException(
        status_code=400,
        detail="Unsupported document type.",
    )


# ============================================================
# FIND NUMBER FROM TEXT
# ============================================================

def find_number(
    text: str,
    patterns: List[str],
) -> Optional[float]:

    for pattern in patterns:

        match = re.search(
            pattern,
            text,
            flags=re.IGNORECASE,
        )

        if not match:
            continue

        try:

            value = match.group(1)

            value = value.replace(
                ",",
                "",
            )

            return float(
                value
            )

        except (
            TypeError,
            ValueError,
        ):

            continue

    return None


# ============================================================
# EXTRACT PROJECT FEATURES FROM DPR
# ============================================================

def extract_project_features(
    text: str,
) -> Dict[str, Optional[float]]:

    lower = text.lower()

    duration_months = find_number(
        lower,
        [
            r"(?:duration|implementation\s+period|implementation\s+duration|completion\s+period)[^\d]{0,100}(\d+(?:\.\d+)?)\s*months?",
            r"(\d+(?:\.\d+)?)\s*months?[^\n]{0,60}(?:duration|implementation|completion)",
        ],
    )

    length_km = find_number(
        lower,
        [
            r"(?:project\s+length|road\s+length)[^\d]{0,80}(\d+(?:\.\d+)?)\s*km",
            r"(?:length)[^\d]{0,50}(\d+(?:\.\d+)?)\s*km",
        ],
    )

    width_m = find_number(
        lower,
        [
            r"(?:carriageway\s+width|road\s+width)[^\d]{0,80}(\d+(?:\.\d+)?)\s*m",
            r"(?:width)[^\d]{0,50}(\d+(?:\.\d+)?)\s*m",
        ],
    )

    population = find_number(
        lower,
        [
            r"(?:population|beneficiaries|beneficiary\s+population)[^\d]{0,100}(\d[\d,]*)",
        ],
    )

    adt = find_number(
        lower,
        [
            r"(?:adt|average\s+daily\s+traffic|traffic\s+volume)[^\d]{0,100}(\d[\d,]*)",
        ],
    )

    structures = find_number(
        lower,
        [
            r"(?:structures|bridges|culverts)[^\d]{0,80}(\d+(?:\.\d+)?)",
        ],
    )

    cbr = find_number(
        lower,
        [
            r"(?:cbr|california\s+bearing\s+ratio)[^\d]{0,80}(\d+(?:\.\d+)?)",
        ],
    )

    right_of_way_m = find_number(
        lower,
        [
            r"(?:right\s+of\s+way|right-of-way|row|r\.o\.w\.)[^\d]{0,80}(\d+(?:\.\d+)?)\s*m",
        ],
    )

    return {
        "duration_months": duration_months,
        "length_km": length_km,
        "width_m": width_m,
        "population": population,
        "adt": adt,
        "structures": structures,
        "cbr": cbr,
        "right_of_way_m": right_of_way_m,
    }


# ============================================================
# DIRECT PROJECT COST EXTRACTION
# ============================================================

def extract_direct_cost(
    text: str,
) -> Optional[float]:

    lower = text.lower()

    return find_number(
        lower,
        [
            r"(?:project\s+cost|capital\s+cost|estimated\s+cost|total\s+project\s+cost)[^\d]{0,100}(\d+(?:\.\d+)?)\s*(?:crore|cr)",

            r"(\d+(?:\.\d+)?)\s*(?:crore|cr)[^\n]{0,80}(?:project\s+cost|capital\s+cost|estimated\s+cost)",
        ],
    )


# ============================================================
# GET DURATION
# ============================================================

def get_project_duration(
    project: Optional[Project],
    extracted_duration: Optional[float],
) -> Optional[float]:

    if (
        extracted_duration is not None
        and extracted_duration > 0
    ):
        return extracted_duration

    if project is not None:

        try:

            project_duration = float(
                getattr(
                    project,
                    "duration_months",
                    0,
                ) or 0
            )

            if project_duration > 0:
                return project_duration

        except (
            TypeError,
            ValueError,
        ):
            pass

    return None


# ============================================================
# UPLOAD DOCUMENT
# ============================================================

@router.post("/upload")
def upload_document(
    project_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> Dict[str, Any]:

    print(
        f"[DOCUMENT] Upload request: "
        f"project_id={project_id}, "
        f"filename={file.filename}"
    )

    # --------------------------------------------------------
    # PROJECT ID
    # --------------------------------------------------------

    if project_id <= 0:

        raise HTTPException(
            status_code=400,
            detail="Invalid project_id.",
        )

    # --------------------------------------------------------
    # CHECK PROJECT
    # --------------------------------------------------------

    project = (
        db.query(Project)
        .filter(
            Project.id == project_id
        )
        .first()
    )

    if project is None:

        raise HTTPException(
            status_code=404,
            detail="Project not found.",
        )

    # --------------------------------------------------------
    # USER
    # --------------------------------------------------------

    uploaded_by = getattr(
        project,
        "user_id",
        None,
    )

    if uploaded_by is None:

        raise HTTPException(
            status_code=500,
            detail=(
                "Project does not have a valid user_id."
            ),
        )

    try:
        uploaded_by = int(
            uploaded_by
        )

    except (
        TypeError,
        ValueError,
    ) as exc:

        raise HTTPException(
            status_code=500,
            detail="Invalid project user_id.",
        ) from exc

    # --------------------------------------------------------
    # FILE
    # --------------------------------------------------------

    if not file.filename:

        raise HTTPException(
            status_code=400,
            detail="No file selected.",
        )

    original_filename = get_safe_filename(
        file.filename
    )

    if not original_filename:

        raise HTTPException(
            status_code=400,
            detail="Invalid file name.",
        )

    extension = get_file_extension(
        original_filename
    )

    if extension not in ALLOWED_EXTENSIONS:

        raise HTTPException(
            status_code=400,
            detail=(
                "Only PDF, DOCX and TXT files are supported."
            ),
        )

    # --------------------------------------------------------
    # PROJECT DIRECTORY
    # --------------------------------------------------------

    project_dir = (
        UPLOAD_DIR /
        str(project_id)
    )

    project_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    # --------------------------------------------------------
    # UNIQUE PATH
    # --------------------------------------------------------

    file_path = (
        project_dir /
        original_filename
    )

    counter = 1

    while file_path.exists():

        file_path = (
            project_dir /
            (
                f"{Path(original_filename).stem}"
                f"_{counter}"
                f"{Path(original_filename).suffix}"
            )
        )

        counter += 1

    # --------------------------------------------------------
    # SAVE FILE
    # --------------------------------------------------------

    try:

        with file_path.open(
            "wb"
        ) as buffer:

            shutil.copyfileobj(
                file.file,
                buffer,
            )

    except Exception as exc:

        print(
            "[DOCUMENT] FILE SAVE ERROR:",
            exc,
        )

        raise HTTPException(
            status_code=500,
            detail=(
                "Could not save uploaded document."
            ),
        ) from exc

    finally:

        try:
            file.file.close()
        except Exception:
            pass

    # --------------------------------------------------------
    # DATABASE RECORD
    # --------------------------------------------------------

    document = Document(
        project_id=project_id,
        filename=original_filename,
        file_path=str(file_path),
        document_type=extension.replace(
            ".",
            "",
        ).upper(),
        uploaded_by=uploaded_by,
    )

    try:

        db.add(
            document
        )

        db.commit()

        db.refresh(
            document
        )

    except Exception as exc:

        db.rollback()

        try:

            if file_path.exists():
                file_path.unlink()

        except Exception as cleanup_error:

            print(
                "[DOCUMENT] FILE CLEANUP ERROR:",
                cleanup_error,
            )

        print(
            "[DOCUMENT] DATABASE ERROR:",
            exc,
        )

        raise HTTPException(
            status_code=500,
            detail=(
                "Could not save document information."
            ),
        ) from exc

    return {
        "message":
            "Document uploaded successfully",

        "id":
            document.id,

        "document_id":
            document.id,

        "filename":
            document.filename,

        "file_path":
            document.file_path,

        "document_type":
            document.document_type,

        "project_id":
            document.project_id,

        "uploaded_by":
            document.uploaded_by,

        "created_at":
            document.created_at,
    }


# ============================================================
# GET PROJECT DOCUMENTS
# ============================================================

@router.get("/project/{project_id}")
def get_project_documents(
    project_id: int,
    db: Session = Depends(get_db),
) -> List[Dict[str, Any]]:

    if project_id <= 0:

        raise HTTPException(
            status_code=400,
            detail="Invalid project_id.",
        )

    project = (
        db.query(Project)
        .filter(
            Project.id == project_id
        )
        .first()
    )

    if project is None:

        raise HTTPException(
            status_code=404,
            detail="Project not found.",
        )

    documents = (
        db.query(Document)
        .filter(
            Document.project_id ==
            project_id
        )
        .order_by(
            Document.created_at.desc()
        )
        .all()
    )

    return [
        {
            "id":
                document.id,

            "filename":
                document.filename,

            "file_path":
                document.file_path,

            "document_type":
                document.document_type,

            "project_id":
                document.project_id,

            "uploaded_by":
                document.uploaded_by,

            "created_at":
                document.created_at,
        }
        for document in documents
    ]


# ============================================================
# GET DOCUMENT TEXT
# ============================================================

@router.get("/{document_id}/text")
def get_document_text(
    document_id: int,
    db: Session = Depends(get_db),
) -> Dict[str, Any]:

    document = (
        db.query(Document)
        .filter(
            Document.id ==
            document_id
        )
        .first()
    )

    if document is None:

        raise HTTPException(
            status_code=404,
            detail="Document not found.",
        )

    if not document.file_path:

        raise HTTPException(
            status_code=404,
            detail="Document file path not available.",
        )

    file_path = Path(
        document.file_path
    )

    if not file_path.exists():

        raise HTTPException(
            status_code=404,
            detail="Document file not found.",
        )

    text = extract_document_text(
        file_path
    )

    return {
        "document_id":
            document.id,

        "filename":
            document.filename,

        "text":
            text,
    }


# ============================================================
# ANALYZE DOCUMENT RISKS
# ============================================================

@router.get("/{document_id}/risks")
def analyze_document_risks(
    document_id: int,
    db: Session = Depends(get_db),
) -> Dict[str, Any]:

    print(
        f"[RISK] Starting analysis "
        f"document_id={document_id}"
    )

    # --------------------------------------------------------
    # GET DOCUMENT
    # --------------------------------------------------------

    document = (
        db.query(Document)
        .filter(
            Document.id ==
            document_id
        )
        .first()
    )

    if document is None:

        raise HTTPException(
            status_code=404,
            detail="Document not found.",
        )

    # --------------------------------------------------------
    # CHECK FILE
    # --------------------------------------------------------

    if not document.file_path:

        raise HTTPException(
            status_code=404,
            detail="Document file path not available.",
        )

    file_path = Path(
        document.file_path
    )

    if not file_path.exists():

        raise HTTPException(
            status_code=404,
            detail="Document file not found.",
        )

    # --------------------------------------------------------
    # EXTRACT TEXT
    # --------------------------------------------------------

    text = extract_document_text(
        file_path
    )

    if not text.strip():

        raise HTTPException(
            status_code=422,
            detail=(
                "The uploaded document does not contain readable text."
            ),
        )

    lower_text = text.lower()

    # --------------------------------------------------------
    # PROJECT
    # --------------------------------------------------------

    project = (
        db.query(Project)
        .filter(
            Project.id ==
            document.project_id
        )
        .first()
    )

    # ========================================================
    # EXTRACT FEATURES
    # ========================================================

    features = extract_project_features(
        text
    )

    print(
        "[RISK] EXTRACTED FEATURES:",
        features,
    )

    # ========================================================
    # DIRECT DPR COST
    # ========================================================

    direct_cost = extract_direct_cost(
        text
    )

    # ========================================================
    # ML COST
    # ========================================================

    capital_cost: Optional[float] = None

    capital_cost_source = "NOT_FOUND"

    capital_cost_source_detail: Optional[str] = None

    prediction_confidence = None

    prediction_model_available = False

    try:

        from app.ml.cost_predictor import (
            get_project_cost_prediction,
            model_available,
        )

        prediction_model_available = (
            model_available()
        )

        if prediction_model_available:

            prediction = (
                get_project_cost_prediction(
                    features
                )
            )

            if prediction is not None:

                capital_cost = round(
                    float(
                        prediction["prediction"]
                    ),
                    2,
                )

                prediction_confidence = (
                    prediction.get(
                        "confidence"
                    )
                )

                capital_cost_source = (
                    "ML_MODEL"
                )

                capital_cost_source_detail = (
                    "Project cost predicted using the trained ML model from features extracted from the uploaded DPR."
                )

    except Exception as exc:

        print(
            "[ML] PREDICTION ERROR:",
            exc,
        )

    # ========================================================
    # DPR COST FALLBACK
    # ========================================================

    if capital_cost is None:

        if direct_cost is not None:

            capital_cost = round(
                direct_cost,
                2,
            )

            capital_cost_source = "DPR"

            capital_cost_source_detail = (
                "Project cost was directly extracted from the uploaded DPR."
            )

    # ========================================================
    # BASELINE FALLBACK
    # ========================================================

    if capital_cost is None:

        baseline = 0.0

        length = features.get(
            "length_km"
        )

        width = features.get(
            "width_m"
        )

        structures = features.get(
            "structures"
        )

        population = features.get(
            "population"
        )

        duration = features.get(
            "duration_months"
        )

        if length is not None:

            baseline += (
                float(length) *
                3.5
            )

        if width is not None:

            baseline += (
                float(width) *
                0.75
            )

        if structures is not None:

            baseline += (
                float(structures) *
                0.80
            )

        if population is not None:

            baseline += (
                float(population) /
                50000.0
            )

        if duration is not None:

            baseline += (
                float(duration) *
                0.10
            )

        if baseline <= 0:

            baseline = 10.0

        capital_cost = round(
            baseline,
            2,
        )

        capital_cost_source = (
            "BASELINE_ESTIMATE"
        )

        capital_cost_source_detail = (
            "No trained ML model was available and no direct DPR cost was found. A feature-based baseline estimate was calculated."
        )

    # ========================================================
    # APPROVED BUDGET
    # ========================================================

    approved_budget: Optional[float] = None

    if project is not None:

        try:

            value = float(
                getattr(
                    project,
                    "approved_budget_cr",
                    0,
                ) or 0
            )

            if value > 0:

                approved_budget = round(
                    value,
                    2,
                )

        except (
            TypeError,
            ValueError,
        ):

            approved_budget = None

    # ========================================================
    # DURATION
    # ========================================================

    duration_months = get_project_duration(
        project,
        features.get(
            "duration_months"
        ),
    )

    # ========================================================
    # RISK LIST
    # ========================================================

    risks: List[
        Dict[str, Any]
    ] = []

    def add_risk(
        category: str,
        severity: str,
        points: float,
        title: str,
        description: str,
        recommendation: str,
        keywords: List[str],
    ) -> None:

        risks.append(
            {
                "category":
                    category,

                "severity":
                    severity,

                "keywords":
                    keywords,

                "points":
                    round(
                        float(points),
                        1,
                    ),

                "title":
                    title,

                "description":
                    description,

                "recommendation":
                    recommendation,
            }
        )

    # ========================================================
    # FINANCIAL RISK
    # ========================================================

    if (
        approved_budget is not None
        and capital_cost is not None
        and approved_budget < capital_cost
    ):

        difference = (
            capital_cost -
            approved_budget
        )

        percentage = (
            difference /
            capital_cost *
            100
        )

        if percentage >= 20:

            severity = "critical"

        elif percentage >= 10:

            severity = "high"

        else:

            severity = "medium"

        points = min(
            95,
            35 + percentage,
        )

        add_risk(
            category="Financial",
            severity=severity,
            points=points,
            title=(
                "Approved budget is below estimated project cost"
            ),
            description=(
                f"Estimated project cost is "
                f"₹{capital_cost:.2f} Cr while "
                f"approved budget is "
                f"₹{approved_budget:.2f} Cr."
            ),
            recommendation=(
                "Reconcile approved budget with "
                "the latest project cost estimate."
            ),
            keywords=[
                "budget",
                "cost",
                "funding",
            ],
        )

    # ========================================================
    # COST INFORMATION
    # ========================================================

    if (
        direct_cost is None
        and not prediction_model_available
    ):

        add_risk(
            category="Financial",
            severity="medium",
            points=52,
            title=(
                "Project cost was not directly stated in the DPR"
            ),
            description=(
                "No directly extractable project "
                "cost was found and the ML model "
                "was not available."
            ),
            recommendation=(
                "Include explicit project cost "
                "and detailed cost breakup."
            ),
            keywords=[
                "project cost",
                "capital cost",
                "budget",
            ],
        )

    # ========================================================
    # SCHEDULE RISK
    # ========================================================

    if duration_months is None:

        add_risk(
            category="Schedule",
            severity="high",
            points=74,
            title="Project duration is missing",
            description=(
                "No reliable implementation "
                "duration was found in the DPR."
            ),
            recommendation=(
                "Add an implementation schedule "
                "with milestones."
            ),
            keywords=[
                "duration",
                "timeline",
                "schedule",
                "milestone",
            ],
        )

    elif duration_months >= 48:

        add_risk(
            category="Schedule",
            severity="high",
            points=72,
            title="Long implementation duration",
            description=(
                f"The DPR indicates an implementation "
                f"period of approximately "
                f"{duration_months:.0f} months."
            ),
            recommendation=(
                "Review milestone dependencies "
                "and provide schedule buffers."
            ),
            keywords=[
                "duration",
                "schedule",
                "timeline",
            ],
        )

    elif duration_months >= 30:

        add_risk(
            category="Schedule",
            severity="medium",
            points=52,
            title="Extended implementation period",
            description=(
                f"The DPR indicates an implementation "
                f"period of approximately "
                f"{duration_months:.0f} months."
            ),
            recommendation=(
                "Monitor project milestones closely."
            ),
            keywords=[
                "duration",
                "schedule",
            ],
        )

    # ========================================================
    # TECHNICAL RISK
    # ========================================================

    technical_points = 0

    if features.get(
        "length_km"
    ) is None:

        technical_points += 18

    if features.get(
        "width_m"
    ) is None:

        technical_points += 15

    if features.get(
        "cbr"
    ) is None:

        technical_points += 18

    if (
        "technical specification"
        not in lower_text
    ):

        technical_points += 15

    if (
        "engineering design"
        not in lower_text
        and "design criteria"
        not in lower_text
    ):

        technical_points += 15

    if technical_points >= 45:

        add_risk(
            category="Technical",
            severity=(
                "high"
                if technical_points >= 65
                else "medium"
            ),
            points=technical_points,
            title="Technical information is incomplete",
            description=(
                "Important engineering and technical "
                "indicators were not clearly identified "
                "from the DPR."
            ),
            recommendation=(
                "Provide engineering design criteria, "
                "technical specifications and assumptions."
            ),
            keywords=[
                "technical",
                "engineering",
                "design",
            ],
        )

    # ========================================================
    # ENVIRONMENTAL RISK
    # ========================================================

    environmental_terms = [
        "environment",
        "environmental",
        "eia",
        "environment impact",
        "clearance",
    ]

    if not any(
        term in lower_text
        for term in environmental_terms
    ):

        add_risk(
            category="Environmental",
            severity="medium",
            points=55,
            title=(
                "Environmental assessment is not evident"
            ),
            description=(
                "Environmental assessment or "
                "clearance information was not "
                "clearly identified."
            ),
            recommendation=(
                "Include applicable environmental "
                "assessment and clearances."
            ),
            keywords=environmental_terms,
        )

    # ========================================================
    # LAND RISK
    # ========================================================

    land_terms = [
        "land acquisition",
        "land requirement",
        "right of way",
        "right-of-way",
    ]

    if not any(
        term in lower_text
        for term in land_terms
    ):

        add_risk(
            category="Land",
            severity="medium",
            points=50,
            title=(
                "Land and right-of-way information is incomplete"
            ),
            description=(
                "Land acquisition or right-of-way "
                "information was not clearly identified."
            ),
            recommendation=(
                "Provide land requirements, right-of-way "
                "details and acquisition status."
            ),
            keywords=land_terms,
        )

    # ========================================================
    # COMPLIANCE RISK
    # ========================================================

    compliance_terms = [
        "approval",
        "regulatory",
        "compliance",
        "statutory",
        "clearance",
    ]

    if not any(
        term in lower_text
        for term in compliance_terms
    ):

        add_risk(
            category="Compliance",
            severity="medium",
            points=58,
            title=(
                "Regulatory compliance information is incomplete"
            ),
            description=(
                "Required approvals, statutory requirements "
                "or regulatory clearances were not clearly "
                "identified."
            ),
            recommendation=(
                "Document all required approvals and "
                "statutory requirements."
            ),
            keywords=compliance_terms,
        )

    # ========================================================
    # PROJECT COMPLEXITY
    # ========================================================

    complexity_score = 0

    length = features.get(
        "length_km"
    )

    population = features.get(
        "population"
    )

    adt = features.get(
        "adt"
    )

    structures = features.get(
        "structures"
    )

    if (
        length is not None
        and length >= 20
    ):

        complexity_score += 12

    if (
        population is not None
        and population >= 50000
    ):

        complexity_score += 10

    if (
        adt is not None
        and adt >= 1000
    ):

        complexity_score += 12

    if (
        structures is not None
        and structures >= 10
    ):

        complexity_score += 12

    if complexity_score >= 20:

        add_risk(
            category="Project Complexity",
            severity=(
                "high"
                if complexity_score >= 30
                else "medium"
            ),
            points=complexity_score,
            title=(
                "Project characteristics indicate elevated complexity"
            ),
            description=(
                "The DPR contains scale, traffic, "
                "population or structural factors "
                "that increase implementation complexity."
            ),
            recommendation=(
                "Increase project monitoring, contingency "
                "planning and implementation controls."
            ),
            keywords=[
                "length",
                "population",
                "traffic",
                "structures",
            ],
        )

    # ========================================================
    # FINAL SCORE
    # ========================================================

    if not risks:

        score = 0

        overall_level = "LOW"

    else:

        score = round(
            min(
                100,
                sum(
                    float(
                        item["points"]
                    )
                    for item in risks
                )
                /
                len(
                    risks
                ),
            ),
            1,
        )

        if score >= 75:

            overall_level = "CRITICAL"

        elif score >= 60:

            overall_level = "HIGH"

        elif score >= 35:

            overall_level = "MEDIUM"

        else:

            overall_level = "LOW"

    # ========================================================
    # RESULT
    # ========================================================

    result = {

        "document_id":
            document.id,

        "project_id":
            document.project_id,

        "filename":
            document.filename,

        "score":
            score,

        "overall_level":
            overall_level,

        "risk_count":
            len(risks),

        "capital_cost_cr":
            capital_cost,

        "capital_cost_source":
            capital_cost_source,

        "capital_cost_source_detail":
            capital_cost_source_detail,

        "prediction_confidence":
            prediction_confidence,

        "prediction_model_available":
            prediction_model_available,

        "approved_budget_cr":
            approved_budget,

        "duration_months":
            duration_months,

        "features":
            features,

        "risks":
            risks,
    }


    print(
        "[RISK] ANALYSIS COMPLETE"
    )

    print(
        result
    )


    return result


# ============================================================
# GET DOCUMENT
# ============================================================

@router.get("/{document_id}")
def get_document(
    document_id: int,
    db: Session = Depends(get_db),
) -> Dict[str, Any]:

    document = (
        db.query(Document)
        .filter(
            Document.id ==
            document_id
        )
        .first()
    )

    if document is None:

        raise HTTPException(
            status_code=404,
            detail="Document not found.",
        )

    return {
        "id":
            document.id,

        "filename":
            document.filename,

        "file_path":
            document.file_path,

        "document_type":
            document.document_type,

        "project_id":
            document.project_id,

        "uploaded_by":
            document.uploaded_by,

        "created_at":
            document.created_at,
    }


# ============================================================
# DELETE DOCUMENT
# ============================================================

@router.delete("/{document_id}")
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
) -> Dict[str, Any]:

    document = (
        db.query(Document)
        .filter(
            Document.id ==
            document_id
        )
        .first()
    )

    if document is None:

        raise HTTPException(
            status_code=404,
            detail="Document not found.",
        )

    if document.file_path:

        file_path = Path(
            document.file_path
        )

        try:

            if file_path.exists():

                file_path.unlink()

        except Exception as exc:

            print(
                "[DOCUMENT] FILE DELETE ERROR:",
                exc,
            )

    try:

        db.delete(
            document
        )

        db.commit()

    except Exception as exc:

        db.rollback()

        raise HTTPException(
            status_code=500,
            detail="Could not delete document.",
        ) from exc

    return {
        "message":
            "Document deleted successfully",

        "document_id":
            document_id,
    }